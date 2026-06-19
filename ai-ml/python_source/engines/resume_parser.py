"""
resume_parser.py
-----------------
Extracts plain text from uploaded resume files (PDF / DOCX) so the
ATS pipeline can score them. Used by POST /ats/upload-resume.
"""

import logging
from io import BytesIO

from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


class ResumeParseError(Exception):
    """Raised when a resume file can't be parsed into usable text."""


def extract_resume_text(file_bytes: bytes, ext: str) -> str:
    ext = ext.lower().lstrip(".")
    if ext == "pdf":
        text = _extract_pdf(file_bytes)
    elif ext == "docx":
        text = _extract_docx(file_bytes)
    else:
        raise ResumeParseError(f"Unsupported file type: .{ext}")

    text = text.strip()
    if not text:
        raise ResumeParseError(
            "No readable text found in this file. It may be a scanned "
            "image with no text layer — try exporting a text-based PDF/DOCX."
        )
    return text


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(file_bytes))
    except Exception as exc:
        logger.warning("PDF parse failed: %s", exc)
        raise ResumeParseError(
            "Could not open this PDF — it may be corrupted or password-protected."
        )

    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as exc:
            logger.warning("PDF page extraction failed: %s", exc)
    return "\n".join(pages)


def _extract_docx(file_bytes: bytes) -> str:
    try:
        doc = Document(BytesIO(file_bytes))
    except Exception as exc:
        logger.warning("DOCX parse failed: %s", exc)
        raise ResumeParseError("Could not open this DOCX — the file may be corrupted.")

    parts = [p.text for p in doc.paragraphs]
    # Resumes often use tables for layout (skills grids, two-column dates) —
    # pull that text too or you silently lose half the resume.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(p for p in parts if p.strip())
